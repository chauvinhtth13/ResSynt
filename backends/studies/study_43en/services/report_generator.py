# backends/studies/study_43en/services/report_generator.py
"""
TMG Report Generator Service - Improved Version

Creates TMG (Trial Management Group) reports in Word format
following OUCRU standard template with proper formatting.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from datetime import datetime
from io import BytesIO
import os
import logging

logger = logging.getLogger(__name__)


class TMGReportGenerator:
    """
    Service class để tạo báo cáo TMG theo format chuẩn OUCRU
    Cải thiện format để khớp với mẫu gốc
    """
    
    # Constants
    DOCUMENT_CODE = 'TM-CTU-003'
    VERSION = '4.0'
    SECTION = 'CTU'
    EFFECTIVE_DATE = '06JAN22'
    
    # Font settings
    FONT_FAMILY = 'Calibri'
    TITLE_SIZE = Pt(20)
    HEADING_SIZE = Pt(12)
    BODY_SIZE = Pt(11)
    SMALL_SIZE = Pt(10)
    TABLE_SIZE = Pt(10)
    
    # Colors
    DARK_BLUE = RGBColor(8, 26, 59)
    HEADER_BG = 'D9E2F3'  # Light blue for header
    BLACK = RGBColor(0, 0, 0)
    
    def __init__(self, study_code: str, reporting_date: datetime):
        self.study_code = study_code
        self.reporting_date = reporting_date
        self.document = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Cấu hình document cơ bản"""
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
    
    def _set_cell_shading(self, cell, color: str):
        """Set background color for cell"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _set_cell_padding(self, cell, padding_pt: int = 3):
        """Set cell padding"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'left', 'bottom', 'right']:
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), str(padding_pt * 20))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)
    
    def _add_header(self, logo_path: str = None):
        """Thêm header với logo, title, version info"""
        header_table = self.document.add_table(rows=2, cols=3)
        header_table.autofit = False
        
        header_table.columns[0].width = Inches(1.5)
        header_table.columns[1].width = Inches(4)
        header_table.columns[2].width = Inches(1.5)
        
        # Logo cell
        logo_cell = header_table.cell(0, 0)
        logo_cell.merge(header_table.cell(1, 0))
        
        if logo_path and os.path.exists(logo_path):
            logo_para = logo_cell.paragraphs[0]
            run = logo_para.add_run()
            try:
                run.add_picture(logo_path, width=Inches(1.2))
            except Exception as e:
                logger.warning(f"Could not add logo: {e}")
                logo_para.add_run("[LOGO]")
        else:
            logo_para = logo_cell.paragraphs[0]
            logo_run = logo_para.add_run("[OUCRU LOGO]")
            logo_run.font.size = self.SMALL_SIZE
            logo_run.font.name = self.FONT_FAMILY
        
        # Title cell
        title_cell = header_table.cell(0, 1)
        title_para = title_cell.paragraphs[0]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("TMG REPORT")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = self.FONT_FAMILY
        
        # Version cell
        version_cell = header_table.cell(0, 2)
        version_para = version_cell.paragraphs[0]
        version_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        version_run = version_para.add_run(f"Version: {self.VERSION}")
        version_run.font.size = self.SMALL_SIZE
        version_run.font.name = self.FONT_FAMILY
        
        # Document info
        doc_cell = header_table.cell(1, 1)
        doc_para = doc_cell.paragraphs[0]
        doc_run = doc_para.add_run(f"Document code: {self.DOCUMENT_CODE}\nSection: {self.SECTION}")
        doc_run.font.size = self.SMALL_SIZE
        doc_run.font.name = self.FONT_FAMILY
        
        # Effective date
        eff_cell = header_table.cell(1, 2)
        eff_para = eff_cell.paragraphs[0]
        eff_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        eff_run = eff_para.add_run(f"Effective: {self.EFFECTIVE_DATE}")
        eff_run.font.size = self.SMALL_SIZE
        eff_run.font.name = self.FONT_FAMILY
        
        self.document.add_paragraph()
        self._add_horizontal_line()
    
    def _add_horizontal_line(self):
        """Thêm đường kẻ ngang"""
        para = self.document.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        para._p.get_or_add_pPr().append(pBdr)
    
    def _add_title_section(self):
        """Thêm title section với study code và date"""
        title_para = self.document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(f"{self.study_code}: UPDATE REPORT")
        title_run.bold = True
        title_run.font.size = Pt(20)
        title_run.font.name = self.FONT_FAMILY
        
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        day = self.reporting_date.day
        month = self.reporting_date.strftime('%b')
        year = self.reporting_date.year
        ordinal = self._get_ordinal_suffix(day)
        
        date_run = date_para.add_run(f"Reporting Date: {month} {day}")
        date_run.bold = True
        date_run.font.size = Pt(14)
        date_run.font.name = self.FONT_FAMILY
        
        sup_run = date_para.add_run(ordinal)
        sup_run.bold = True
        sup_run.font.size = Pt(14)
        sup_run.font.name = self.FONT_FAMILY
        sup_run.font.superscript = True
        
        year_run = date_para.add_run(f", {year}")
        year_run.bold = True
        year_run.font.size = Pt(14)
        year_run.font.name = self.FONT_FAMILY
        
        self.document.add_paragraph()
    
    def _get_ordinal_suffix(self, day: int) -> str:
        if 11 <= day <= 13:
            return 'th'
        return {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')
    
    def _add_section_heading(self, number: int, title: str, underline: bool = True):
        """Thêm heading cho section"""
        para = self.document.add_paragraph()
        
        num_run = para.add_run(f"{number}. ")
        num_run.bold = True
        num_run.font.size = self.HEADING_SIZE
        num_run.font.name = self.FONT_FAMILY
        
        title_run = para.add_run(title)
        title_run.bold = True
        title_run.font.size = self.HEADING_SIZE
        title_run.font.name = self.FONT_FAMILY
        if underline:
            title_run.underline = True
    
    def _add_action_points_table(self, action_points: list):
        """Thêm bảng Outstanding Action Points"""
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        headers = ['Action Point', 'Actioned by', 'Complete?']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            self._set_cell_shading(cell, self.HEADER_BG)
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.size = self.TABLE_SIZE
            run.font.name = self.FONT_FAMILY
        
        table.columns[0].width = Inches(4)
        table.columns[1].width = Inches(1.5)
        table.columns[2].width = Inches(1)
        
        if action_points:
            for item in action_points:
                row = table.add_row().cells
                
                action_run = row[0].paragraphs[0].add_run(f"• {item.get('action', '')}")
                action_run.font.size = self.TABLE_SIZE
                action_run.font.name = self.FONT_FAMILY
                
                by_run = row[1].paragraphs[0].add_run(item.get('actioned_by', ''))
                by_run.font.size = self.TABLE_SIZE
                by_run.font.name = self.FONT_FAMILY
                
                complete_run = row[2].paragraphs[0].add_run(item.get('complete', ''))
                complete_run.font.size = self.TABLE_SIZE
                complete_run.font.name = self.FONT_FAMILY
        else:
            row = table.add_row().cells
            run = row[0].paragraphs[0].add_run("No outstanding action points")
            run.font.size = self.TABLE_SIZE
            run.font.italic = True
        
        self.document.add_paragraph()
    
    def _add_recruitment_table(self, recruitment_data: dict):
        """Thêm bảng thống kê recruitment - format mới với Patient và Contact"""
        table = self.document.add_table(rows=3, cols=3)
        table.style = 'Table Grid'
        
        # Header row
        headers = ['Category', 'Patients', 'Contacts']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            self._set_cell_shading(cell, self.HEADER_BG)
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.size = self.TABLE_SIZE
            run.font.name = self.FONT_FAMILY
        
        rows_data = [
            ('Total Screened', 
             recruitment_data.get('total_screened_patients', 0),
             recruitment_data.get('total_screened_contacts', 0)),
            ('Total Enrolled', 
             recruitment_data.get('total_enrolled_patients', 0),
             recruitment_data.get('total_enrolled_contacts', 0)),
        ]
        
        for i, (label, patients, contacts) in enumerate(rows_data):
            row = table.rows[i + 1]
            
            label_run = row.cells[0].paragraphs[0].add_run(label)
            label_run.bold = True
            label_run.font.size = self.TABLE_SIZE
            label_run.font.name = self.FONT_FAMILY
            
            pat_run = row.cells[1].paragraphs[0].add_run(str(patients))
            pat_run.font.size = self.TABLE_SIZE
            pat_run.font.name = self.FONT_FAMILY
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            con_run = row.cells[2].paragraphs[0].add_run(str(contacts))
            con_run.font.size = self.TABLE_SIZE
            con_run.font.name = self.FONT_FAMILY
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph()
    
    def _add_sample_processing_table(self, sample_data: dict):
        """
        Thêm bảng Sample Processing theo format dashboard
        Schedule | Patient Total | Patient Blood | Contact Total | Contact Blood
        """
        table = self.document.add_table(rows=5, cols=5)
        table.style = 'Table Grid'
        
        # Header row
        headers = ['Schedule', 'Patient Total', 'Patient Blood', 'Contact Total', 'Contact Blood']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            self._set_cell_shading(cell, self.HEADER_BG)
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.size = self.TABLE_SIZE
            run.font.name = self.FONT_FAMILY
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        patient_stats = sample_data.get('patient', {})
        contact_stats = sample_data.get('contact', {})
        
        schedules = [
            ('Day 1 (Enrollment)', 'visit1'),
            ('Day 10 (±3 days)', 'visit2'),
            ('Day 28 (±3 days)', 'visit3'),
            ('Day 90 (±3 days)', 'visit4'),
        ]
        
        for row_idx, (schedule_name, visit_key) in enumerate(schedules):
            row = table.rows[row_idx + 1]
            
            # Schedule name
            run = row.cells[0].paragraphs[0].add_run(schedule_name)
            run.font.size = self.TABLE_SIZE
            run.font.name = self.FONT_FAMILY
            
            # Patient total
            p_data = patient_stats.get(visit_key, {})
            p_total = p_data.get('total', 0) if p_data else 0
            run = row.cells[1].paragraphs[0].add_run(str(p_total))
            run.font.size = self.TABLE_SIZE
            run.font.name = self.FONT_FAMILY
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Patient blood
            p_blood = p_data.get('blood', 0) if p_data else 0
            run = row.cells[2].paragraphs[0].add_run(str(p_blood))
            run.font.size = self.TABLE_SIZE
            run.font.name = self.FONT_FAMILY
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact total
            c_data = contact_stats.get(visit_key, {})
            c_total = c_data.get('total', 0) if c_data else 0
            # Day 10 is N/A for contacts
            if visit_key == 'visit2':
                run = row.cells[3].paragraphs[0].add_run('N/A')
                run.font.italic = True
            else:
                run = row.cells[3].paragraphs[0].add_run(str(c_total))
            run.font.size = self.TABLE_SIZE
            run.font.name = self.FONT_FAMILY
            row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Contact blood
            c_blood = c_data.get('blood', 0) if c_data else 0
            if visit_key == 'visit2':
                run = row.cells[4].paragraphs[0].add_run('N/A')
                run.font.italic = True
            else:
                run = row.cells[4].paragraphs[0].add_run(str(c_blood))
            run.font.size = self.TABLE_SIZE
            run.font.name = self.FONT_FAMILY
            row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph()
    
    def _add_deviations_table(self, deviations: list):
        """Thêm bảng Protocol Deviations"""
        table = self.document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        headers = ['Subject ID', 'Deviation description', 'Violation?', 'Action']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            self._set_cell_shading(cell, self.HEADER_BG)
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.size = self.TABLE_SIZE
            run.font.name = self.FONT_FAMILY
        
        if deviations:
            for item in deviations:
                row = table.add_row().cells
                for col, key in enumerate(['subject_id', 'deviation', 'violation', 'action']):
                    run = row[col].paragraphs[0].add_run(item.get(key, ''))
                    run.font.size = self.TABLE_SIZE
                    run.font.name = self.FONT_FAMILY
        else:
            row = table.add_row().cells
            run = row[0].paragraphs[0].add_run("No protocol deviations reported")
            run.font.size = self.TABLE_SIZE
            run.font.italic = True
        
        self.document.add_paragraph()
    
    def _add_safety_table(self, safety_data: dict):
        """Thêm bảng Safety Reporting"""
        table = self.document.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        # Header
        self._set_cell_shading(table.rows[0].cells[0], self.HEADER_BG)
        self._set_cell_shading(table.rows[0].cells[1], self.HEADER_BG)
        table.rows[0].cells[0].paragraphs[0].add_run("Category").bold = True
        table.rows[0].cells[1].paragraphs[0].add_run("Count").bold = True
        
        rows_data = [
            ('Total AEs reported', safety_data.get('total_ae', 0)),
            ('Total SAEs reported', safety_data.get('total_sae', 0)),
            ('Deaths', safety_data.get('deaths', 0)),
            ('AEs leading to discontinuation', safety_data.get('ae_discontinuation', 0)),
        ]
        
        for i, (label, value) in enumerate(rows_data):
            row = table.rows[i + 1]
            
            label_run = row.cells[0].paragraphs[0].add_run(label)
            label_run.font.size = self.TABLE_SIZE
            label_run.font.name = self.FONT_FAMILY
            
            value_run = row.cells[1].paragraphs[0].add_run(str(value))
            value_run.font.size = self.TABLE_SIZE
            value_run.font.name = self.FONT_FAMILY
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph()
    
    def _add_text_content(self, content: str):
        """Thêm nội dung text"""
        para = self.document.add_paragraph()
        if content:
            run = para.add_run(content)
        else:
            run = para.add_run("N/A")
            run.italic = True
        run.font.size = self.BODY_SIZE
        run.font.name = self.FONT_FAMILY
        
        self.document.add_paragraph()
    
    def generate(self, report_data: dict, logo_path: str = None) -> BytesIO:
        """Tạo báo cáo hoàn chỉnh"""
        
        # Header
        self._add_header(logo_path)
        
        # Title
        self._add_title_section()
        
        # Section 1: Outstanding Action Points
        self._add_section_heading(1, "Outstanding Action Points")
        self._add_action_points_table(report_data.get('action_points', []))
        
        # Section 2: General Trial Procedures
        self._add_section_heading(2, "General Trial Procedures")
        self._add_text_content(report_data.get('general_procedures', ''))
        
        # Section 3: Ethics & Regulatory
        self._add_section_heading(3, "Ethics & Regulatory")
        self._add_text_content(report_data.get('ethics_regulatory', ''))
        
        # Section 4: Study Amendments
        self._add_section_heading(4, "Study Amendments")
        self._add_text_content(report_data.get('study_amendments', ''))
        
        # Section 5: Recruitment
        self._add_section_heading(5, "Recruitment")
        self._add_recruitment_table(report_data.get('recruitment', {}))
        
        # Section 6: Protocol Deviations/Violations
        self._add_section_heading(6, "Protocol Deviations/Violations")
        self._add_deviations_table(report_data.get('deviations', []))
        
        # Section 7: Sample Processing
        self._add_section_heading(7, "Sample Processing")
        sample_data = report_data.get('sample_processing', {})
        if sample_data and (sample_data.get('patient') or sample_data.get('contact')):
            self._add_sample_processing_table(sample_data)
        else:
            self._add_text_content('')
        
        # Section 8: Data Management
        self._add_section_heading(8, "Data Management")
        self._add_text_content(report_data.get('data_management', ''))
        
        # Section 9: Safety Reporting
        self._add_section_heading(9, "Safety Reporting")
        safety_data = report_data.get('safety_reporting', {})
        if safety_data:
            self._add_safety_table(safety_data)
        else:
            self._add_text_content('')
        
        # Section 10: AOB
        self._add_section_heading(10, "AOB")
        self._add_text_content(report_data.get('aob', ''))
        
        # Save to BytesIO
        buffer = BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer
